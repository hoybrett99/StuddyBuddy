# tests/test_document_service.py

import pytest
from pathlib import Path
from backend.app.services.document_services import DocumentService
from backend.app.models import DocumentMetaData, DocumentChunk, FileType
from backend.app.config import Settings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import tempfile
import shutil

# ============================================================================
# Fixtures - Reusable test data
# ============================================================================

@pytest.fixture
def temp_upload_dir(tmp_path):
    """Create a temporary upload directory for testing"""
    upload_dir = tmp_path / "uploads"
    upload_dir.mkdir()
    return upload_dir


@pytest.fixture
def document_service(temp_upload_dir, monkeypatch):
    """Create a DocumentService with a temporary upload directory"""
    # Mock the settings to use our temp directory
    def mock_get_settings():
        return Settings(
            google_api_key="test_key",
            claude_api_key="test_key",
            upload_dir=str(temp_upload_dir)
        )
    
    monkeypatch.setattr("backend.app.services.document_services.get_settings", mock_get_settings)
    return DocumentService()


@pytest.fixture
def sample_metadata():
    """Create sample document metadata"""
    return DocumentMetaData(
        filename="test_document.pdf",
        file_type=FileType.PDF,
        file_size_bytes=1024
    )


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF file for testing"""
    from pypdf import PdfWriter
    
    pdf_path = tmp_path / "test.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    
    with open(pdf_path, 'wb') as f:
        writer.write(f)
    
    return pdf_path


@pytest.fixture
def sample_txt(tmp_path):
    """Create a sample TXT file for testing"""
    txt_path = tmp_path / "test.txt"
    txt_path.write_text("This is a test text file.\nIt has multiple lines.\n", encoding='utf-8')
    return txt_path


# ============================================================================
# Test Initialization
# ============================================================================

def test_document_service_creation(document_service):
    """Test that DocumentService can be created"""
    assert document_service is not None
    assert isinstance(document_service, DocumentService)


def test_document_service_has_settings(document_service):
    """Test that DocumentService loads settings"""
    assert document_service.settings is not None
    assert isinstance(document_service.settings, Settings)


def test_document_service_has_upload_dir(document_service):
    """Test that DocumentService has upload directory"""
    assert document_service.upload_dir is not None
    assert isinstance(document_service.upload_dir, Path)
    assert document_service.upload_dir.exists()


def test_document_service_has_text_splitter(document_service):
    """Test that DocumentService has text splitter"""
    assert document_service.text_splitter is not None
    assert isinstance(document_service.text_splitter, RecursiveCharacterTextSplitter)


def test_text_splitter_configuration(document_service):
    """Test that text splitter uses correct configuration"""
    assert document_service.text_splitter._chunk_size == document_service.settings.chunk_size
    assert document_service.text_splitter._chunk_overlap == document_service.settings.chunk_overlap


# ============================================================================
# Test save_file
# ============================================================================

@pytest.mark.asyncio
async def test_save_file(document_service):
    """Test saving a file"""
    filename = "test.pdf"
    content = b"PDF content here"
    
    saved_path = await document_service.save_file(filename, content)
    
    # Check that file was saved
    assert saved_path.exists()
    assert saved_path.is_file()
    
    # Check content
    with open(saved_path, 'rb') as f:
        assert f.read() == content


@pytest.mark.asyncio
async def test_save_file_creates_unique_filename(document_service):
    """Test that save_file creates unique filenames"""
    filename = "document.pdf"
    content1 = b"First file"
    content2 = b"Second file"
    
    path1 = await document_service.save_file(filename, content1)
    path2 = await document_service.save_file(filename, content2)
    
    # Filenames should be different
    assert path1 != path2
    assert path1.name != path2.name
    
    # Both should contain original filename
    assert filename in path1.name
    assert filename in path2.name


@pytest.mark.asyncio
async def test_save_file_in_upload_dir(document_service):
    """Test that files are saved in the upload directory"""
    filename = "test.pdf"
    content = b"Test content"
    
    saved_path = await document_service.save_file(filename, content)
    
    # Should be in upload directory
    assert saved_path.parent == document_service.upload_dir


# ============================================================================
# Test _extract_txt
# ============================================================================

def test_extract_txt(document_service, sample_txt):
    """Test extracting text from TXT file"""
    text = document_service._extract_txt(sample_txt)
    
    assert text is not None
    assert "test text file" in text
    assert "multiple lines" in text


def test_extract_txt_preserves_content(document_service, tmp_path):
    """Test that TXT extraction preserves exact content"""
    expected_text = "Line 1\nLine 2\nLine 3"
    txt_path = tmp_path / "test.txt"
    txt_path.write_text(expected_text, encoding='utf-8')
    
    extracted_text = document_service._extract_txt(txt_path)
    
    assert extracted_text == expected_text


# ============================================================================
# Test _extract_pdf
# ============================================================================

def test_extract_pdf_returns_string(document_service, sample_pdf):
    """Test that PDF extraction returns a string"""
    text = document_service._extract_pdf(sample_pdf)
    
    assert isinstance(text, str)


def test_extract_pdf_with_text(document_service, tmp_path):
    """Test extracting text from PDF with actual content"""
    # Note: Creating PDFs with text is complex, so this is a basic test
    from pypdf import PdfWriter
    
    pdf_path = tmp_path / "text.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    
    with open(pdf_path, 'wb') as f:
        writer.write(f)
    
    text = document_service._extract_pdf(pdf_path)
    assert isinstance(text, str)


# ============================================================================
# Test _extract_docx
# ============================================================================

def test_extract_docx(document_service, tmp_path):
    """Test extracting text from DOCX file"""
    from docx import Document
    
    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(str(docx_path))
    
    text = document_service._extract_docx(docx_path)
    
    assert "First paragraph" in text
    assert "Second paragraph" in text
    assert "\n\n" in text  # Paragraphs separated by double newline


def test_extract_docx_empty_document(document_service, tmp_path):
    """Test extracting from empty DOCX"""
    from docx import Document
    
    docx_path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(docx_path))
    
    text = document_service._extract_docx(docx_path)
    
    assert text == ""


# ============================================================================
# Test text_extraction (router method)
# ============================================================================

def test_text_extraction_pdf(document_service, sample_pdf):
    """Test text extraction routing for PDF"""
    text = document_service.text_extraction(sample_pdf, FileType.PDF)
    
    assert isinstance(text, str)


def test_text_extraction_txt(document_service, sample_txt):
    """Test text extraction routing for TXT"""
    text = document_service.text_extraction(sample_txt, FileType.TXT)
    
    assert "test text file" in text


def test_text_extraction_docx(document_service, tmp_path):
    """Test text extraction routing for DOCX"""
    from docx import Document
    
    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Test content")
    doc.save(str(docx_path))
    
    text = document_service.text_extraction(docx_path, FileType.DOCX)
    
    assert "Test content" in text


def test_text_extraction_unsupported_type(document_service, sample_txt):
    """Test that unsupported file types raise ValueError"""
    with pytest.raises(ValueError) as exc_info:
        # Create a fake FileType (this won't work in practice, but tests the else clause)
        document_service.text_extraction(sample_txt, "INVALID")
    
    assert "Unsupported file type" in str(exc_info.value)


# ============================================================================
# Test create_chunks
# ============================================================================

def test_create_chunks_basic(document_service, sample_metadata):
    """Test creating chunks from text"""
    text = "This is a test document. " * 100  # Create longer text
    document_id = "doc_123"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    assert len(chunks) > 0
    assert all(isinstance(chunk, DocumentChunk) for chunk in chunks)


def test_create_chunks_ids(document_service, sample_metadata):
    """Test that chunk IDs are correctly generated"""
    text = "Test text. " * 200
    document_id = "doc_abc"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    for idx, chunk in enumerate(chunks):
        assert chunk.chunk_id == f"{document_id}_chunk_{idx}"
        assert chunk.document_id == document_id
        assert chunk.chunk_index == idx


def test_create_chunks_positions(document_service, sample_metadata):
    """Test that chunk positions are tracked correctly"""
    text = "A" * 3000  # 3000 characters
    document_id = "doc_123"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    # First chunk should start at 0
    assert chunks[0].start_char == 0
    
    # Each chunk's end should be after its start
    for chunk in chunks:
        assert chunk.end_char > chunk.start_char
    
    # Positions should be sequential
    for i in range(len(chunks) - 1):
        # Next chunk starts where or near where previous ended (due to overlap)
        assert chunks[i + 1].start_char >= chunks[i].start_char


def test_create_chunks_metadata(document_service, sample_metadata):
    """Test that metadata is attached to chunks"""
    text = "Test text"
    document_id = "doc_123"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    for chunk in chunks:
        assert chunk.metadata == sample_metadata
        assert chunk.metadata.filename == "test_document.pdf"


def test_create_chunks_short_text(document_service, sample_metadata):
    """Test creating chunks from short text"""
    text = "Short text"
    document_id = "doc_123"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    # Short text should create one chunk
    assert len(chunks) == 1
    assert chunks[0].text == text


def test_create_chunks_long_text(document_service, sample_metadata):
    """Test creating chunks from long text"""
    text = "A" * 5000  # Text longer than chunk_size (1000)
    document_id = "doc_123"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    # Should create multiple chunks
    assert len(chunks) > 1


def test_create_chunks_text_content(document_service, sample_metadata):
    """Test that chunk text is correctly extracted"""
    text = "This is chunk 1. " * 100 + "This is chunk 2. " * 100
    document_id = "doc_123"
    
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    # Each chunk should have non-empty text
    for chunk in chunks:
        assert chunk.text
        assert len(chunk.text) > 0


# ============================================================================
# Integration Tests
# ============================================================================

@pytest.mark.asyncio
async def test_full_document_processing_flow(document_service, sample_metadata, tmp_path):
    """Test the complete flow: save -> extract -> chunk"""
    # 1. Create a text file
    content = b"This is a test document. " * 200  # Long content
    filename = "test.txt"
    
    # 2. Save file
    file_path = await document_service.save_file(filename, content)
    assert file_path.exists()
    
    # 3. Extract text
    text = document_service.text_extraction(file_path, FileType.TXT)
    assert len(text) > 0
    
    # 4. Create chunks
    document_id = "doc_integration_test"
    chunks = document_service.create_chunks(text, document_id, sample_metadata)
    
    # Verify chunks
    assert len(chunks) > 0
    assert chunks[0].document_id == document_id
    assert chunks[0].metadata == sample_metadata