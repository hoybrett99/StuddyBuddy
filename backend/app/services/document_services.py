"""
Service for handling document operations.
"""
import uuid
from pathlib import Path
import re
from typing import List
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.models import (
    DocumentMetaData,
    DocumentChunk,
    FileType
)
from app.config import get_settings

class DocumentService:
    """
    A class groups related functions (methods) together.
    
    'self' refers to the instance of the class.
    Think of self as "this specific DocumentService object"
    """
    def __init__(self):
        self.settings = get_settings()
        self.upload_dir = Path(self.settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        # Text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = self.settings.chunk_size,
            chunk_overlap = self.settings.chunk_overlap,
            length_function = len
        )

    # Async functions
    async def save_file(self, filename: str, content: bytes) -> Path:
                """
                Save uploaded file to disk.
                
                Args:
                    filename: Name of the file
                    content: The file's bytes
                    
                Returns:
                    Path: Where the file was saved
                """
                # Create unique filenames to avoid conflicts
                file_id = uuid.uuid4().hex[:8]
                safe_filename = f"{file_id}_{filename}"
                file_path = self.upload_dir / safe_filename

                with open(file_path, 'wb') as f:
                        f.write(content)

                        return file_path
    
    def extract_text(self, file_path: Path, file_type: FileType) -> str:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the file
            file_type: Type of file (PDF, TXT, DOCX)
            
        Returns:
            str: Extracted text
            
        Raises:
            ValueError: If file type not supported
        """
        print(f"Extracting text from {file_type.value.upper()} file...")
        
        if file_type == FileType.PDF:
            return self._extract_pdf(file_path)
        elif file_type == FileType.TXT:
            return self._extract_txt(file_path)
        elif file_type == FileType.DOCX:
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _clean_text(self, text: str) -> str:
        """
        Minimal text cleaning - only fixes obvious OCR errors.
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Lightly cleaned text
        """
        # Fix smart quotes and dashes
        replacements = {
            ''': "'",
            ''': "'",
            '"': '"',
            '"': '"',
            '—': '-',
            '–': '-',
            '…': '...',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Replace multiple spaces with single space (but preserve newlines)
        text = re.sub(r' +', ' ', text)
        
        # Replace more than 3 newlines with 2 newlines
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        
        # Remove trailing/leading whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()

    def _extract_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF - optimized for large files.
        
        Tries:
        1. pdfplumber (best for digital PDFs)
        2. pypdf (fallback)
        
        Returns the cleanest result.
        """
        import time
        start_time = time.time()
        
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        print(f"PDF size: {file_size_mb:.1f}MB")
        
        results = []
        
        # Method 1: pdfplumber (slower but better quality)
        try:
            import pdfplumber
            print("Using pdfplumber for PDF extraction")
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"Processing {total_pages} pages...")
                
                for i, page in enumerate(pdf.pages, 1):
                    # Progress indicator for large files
                    if i % 50 == 0 or i == total_pages:
                        elapsed = time.time() - start_time
                        print(f"  Page {i}/{total_pages} ({elapsed:.1f}s elapsed)")
                    
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            text = "\n\n".join(text_parts)
            
            if text:
                cleaned = self._clean_text(text)
                results.append(('pdfplumber', cleaned))
                print(f"✓ Extracted {len(cleaned):,} characters in {time.time() - start_time:.1f}s")
        
        except Exception as e:
            print(f"⚠️ pdfplumber failed: {e}")
        
        # Method 2: pypdf (faster, lower quality)
        if not results:
            try:
                from pypdf import PdfReader
                print("Trying pypdf as fallback...")
                
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
                print(f"Processing {total_pages} pages...")
                
                text_parts = []
                for i, page in enumerate(reader.pages, 1):
                    if i % 50 == 0 or i == total_pages:
                        print(f"  Page {i}/{total_pages}")
                    
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                text = "\n\n".join(text_parts)
                
                if text:
                    cleaned = self._clean_text(text)
                    results.append(('pypdf', cleaned))
                    print(f"Extracted {len(cleaned):,} characters in {time.time() - start_time:.1f}s")
            
            except Exception as e:
                print(f"pypdf failed: {e}")
        
        # Return best result
        if results:
            best = max(results, key=lambda x: len(x[1]))
            return best[1]
        else:
            raise ValueError("Could not extract text from PDF with any method")

    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"Read TXT file with {encoding} encoding")
                return self._clean_text(text)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode TXT file with any encoding")

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            import docx
            doc = docx.Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n\n".join(text_parts)
            print(f"Extracted {len(text):,} characters from DOCX")
            
            return self._clean_text(text)
        
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
        
    def create_chunks(
        self,
        text: str,
        document_id: str,
        metadata: dict
    ) -> List[DocumentChunk]:
        """
        Create optimized chunks for general PDFs.
        
        Optimizations:
        - 800 char chunks (better than 1000 for dense content)
        - Comprehensive separator list (preserves sentences)
        - Position tracking (for debugging/analysis)
        """
        
        # Configure splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=150,
            length_function=len,
            separators=[
                "\n\n\n",  # Major breaks (3+ newlines)
                "\n\n",    # Paragraph breaks
                "\n",      # Line breaks
                ". ",      # Sentence ends (period + space)
                "! ",      # Exclamations
                "? ",      # Questions
                "; ",      # Semicolons
                ": ",      # Colons
                ", ",      # Commas
                " ",       # Spaces
                "",        # Characters (last resort)
            ],
            is_separator_regex=False,
        )
        
        # Split text
        raw_chunks = text_splitter.split_text(text)
        
        # Create chunks with metadata
        chunks = []
        current_pos = 0
        
        for i, chunk_text in enumerate(raw_chunks):
            # Clean whitespace
            cleaned = chunk_text.strip()
            if not cleaned:
                continue
            
            # Track position in original text
            start_char = text.find(chunk_text, current_pos)
            if start_char == -1:
                start_char = current_pos
            end_char = start_char + len(chunk_text)
            
            # Create chunk
            chunk = DocumentChunk(
                chunk_id=str(uuid.uuid4()),
                document_id=document_id,
                text=cleaned,
                chunk_index=len(chunks),  # Use actual index (skipping empty chunks)
                start_char=start_char,
                end_char=end_char,
                metadata={
                    **metadata,
                    "chunk_index": len(chunks),
                    "total_chunks": len(raw_chunks),
                    "chunk_size": len(cleaned),
                }
            )
            chunks.append(chunk)
            current_pos = end_char
        
        print(f"Created {len(chunks)} chunks (avg {sum(len(c.text) for c in chunks) / len(chunks):.0f} chars)")
        return chunks

if __name__ == "__main__":
    """
    Manual testing with sample documents.
    Run: python -m backend.app.services.document_services
    """
    try:
        print("=" * 60)
        print("DocumentService Manual Test with Sample Documents")
        print("=" * 60)
        
        # Initialize service
        doc1 = DocumentService()
        print("✓ DocumentService initialized\n")
        
        # ==============================================================
        # Test 1: Create chunks from sample text
        # ==============================================================
        print("Test 1: Text Chunking")
        print("-" * 60)
        
        # Sample document text
        sample_text = """
        Introduction to Biology
        
        Biology is the natural science that studies life and living organisms.
        This includes their physical structure, chemical processes, molecular 
        interactions, physiological mechanisms, development and evolution.
        
        Chapter 1: Cell Structure
        
        Cells are the basic building blocks of all living things. The human body 
        is composed of trillions of cells. They provide structure for the body, 
        take in nutrients from food, convert those nutrients into energy, and 
        carry out specialized functions.
        
        Chapter 2: Genetics
        
        Genetics is a branch of biology concerned with the study of genes, 
        genetic variation, and heredity in organisms. Gregor Mendel, a scientist 
        and Augustinian friar, gained posthumous recognition as the founder of 
        the modern science of genetics.
        """ * 10  # Repeat to make it longer for multiple chunks
        
        # Create metadata
        from app.models import DocumentMetaData, FileType
        
        metadata = DocumentMetaData(
            filename="biology_textbook.pdf",
            file_type=FileType.PDF,
            file_size_bytes=len(sample_text.encode('utf-8'))
        )
        
        # Create chunks
        document_id = "doc_biology_001"
        chunks = doc1.create_chunks(sample_text, document_id, metadata)
        
        print(f"✓ Created {len(chunks)} chunks")
        print(f"✓ First chunk ID: {chunks[0].chunk_id}")
        print(f"✓ First chunk preview: {chunks[0].text[:100]}...")
        print(f"✓ Chunk positions: start={chunks[0].start_char}, end={chunks[0].end_char}")
        
        # ==============================================================
        # Test 2: Display all chunks summary
        # ==============================================================
        print("\nTest 2: Chunks Summary")
        print("-" * 60)
        
        for i, chunk in enumerate(chunks):
            print(f"Chunk {i}:")
            print(f"  ID: {chunk.chunk_id}")
            print(f"  Length: {len(chunk.text)} chars")
            print(f"  Position: {chunk.start_char} -> {chunk.end_char}")
            print(f"  Preview: {chunk.text[:80].strip()}...")
            print()
        
        print("=" * 60)
        print("All tests completed successfully! ✓")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()

        